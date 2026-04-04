"""
文档生成：把数据库中的题目生成 Word 文档。
"""
import os
import json
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from shuati.core.database import get_threads_by_date_range, get_questions_by_thread, get_blocks_by_thread
from shuati.core.config import DATA_DIR
from shuati.core.question_parser import _extract_leading_seq
from PIL import Image, ImageChops
import tempfile

LARGE_IMAGE_TRIGGER_WIDTH = Cm(12)
LARGE_IMAGE_TRIGGER_HEIGHT = Cm(10)
EXPORT_IMAGE_MAX_WIDTH = Cm(10.5)
EXPORT_IMAGE_MAX_HEIGHT = Cm(8.5)
DIAGRAM_IMAGE_MAX_WIDTH = Cm(9.5)
DIAGRAM_IMAGE_MAX_HEIGHT = Cm(5.8)


def _downscale_to_max(shape, max_width, max_height):
    ratio = min(float(max_width) / float(shape.width), float(max_height) / float(shape.height), 1.0)
    if ratio < 1.0:
        shape.width = int(float(shape.width) * ratio)
        shape.height = int(float(shape.height) * ratio)


def _is_placeholder_only(content: str) -> bool:
    return bool(re.match(r"^[\u200B-\u200D\uFEFF\s]*\d+[、\.\)]?\s*$", str(content or "")))


def generate_word(start_date: str, end_date: str, output_path: str = None, source_thread_ids: list[str] | None = None, preprocess_images: bool = True) -> str:
    threads = get_threads_by_date_range(start_date, end_date)
    if source_thread_ids:
        picked = set(source_thread_ids)
        threads = [t for t in threads if t["thread_id"] in picked]
    if not threads:
        raise ValueError(f"在 {start_date} 到 {end_date} 之间没有题目")

    doc = Document()
    _setup_page(doc)

    doc.add_paragraph()

    for thread in threads:
        source_title = doc.add_paragraph(style='Heading 1')
        source_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        source_title_run = source_title.add_run(f"{thread['subject']}")
        source_title_run.bold = True
        try:
            source_title_run.font.size = Pt(20)
            source_title_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        except Exception:
            pass
        doc.add_paragraph()

        questions = get_questions_by_thread(thread["thread_id"])
        blocks = get_blocks_by_thread(thread["thread_id"])

        image_marks = {}
        seq_has_text = {}
        for q in questions:
            seq = q.get("seq", 0)
            content = str(q.get("content") or "").strip()
            seq_has_text[seq] = bool(content and not _is_placeholder_only(content))
            for img in json.loads(q.get("images") or "[]"):
                if img:
                    key = os.path.normpath(os.path.abspath(img))
                    if key not in image_marks:
                        image_marks[key] = (seq, "配图")
            for ans in json.loads(q.get("answers") or "[]"):
                if ans:
                    key = os.path.normpath(os.path.abspath(ans))
                    image_marks[key] = (seq, "答案")

        in_answer_section = False
        text_q_seen = 0
        
        for block in blocks:
            ctype = block.get("content_type")
            if ctype == 11:
                text = (block.get("text") or "").strip()
                if text:
                    if any(k in text for k in ["解答", "答案", "答：", "答:"]):
                        in_answer_section = True
                    first_line = next((s.strip() for s in text.splitlines() if s.strip()), "")
                    first_seq = _extract_leading_seq(first_line)
                    if first_seq:
                        text_q_seen = max(text_q_seen, first_seq)
                        in_answer_section = False
                    p = doc.add_paragraph(text)
                    p.paragraph_format.space_after = Pt(6)
            elif ctype == 4:
                local = block.get("image_local")
                key = os.path.normpath(os.path.abspath(local)) if local else ""
                seq, mark = image_marks.get(key, (0, "图片"))
                if key not in image_marks:
                    continue
                answer_by_mark = mark == "答案"
                if answer_by_mark or in_answer_section:
                    in_answer_section = True
                    continue
                if local and os.path.exists(local):
                    image_to_insert = local
                    temp_files_to_clean = []
                    is_diagram = mark == "配图"
                    
                    cropped_dims = None
                    if preprocess_images and seq > 0 and not is_diagram:
                        try:
                            im = Image.open(local).convert("RGB")
                            bg = Image.new(im.mode, im.size, (255, 255, 255))
                            diff = ImageChops.difference(im, bg)
                            diff = ImageChops.add(diff, diff, 2.0, -100)
                            bbox = diff.getbbox()
                            if bbox:
                                padding = 10
                                padded_bbox = (
                                    max(0, bbox[0] - padding),
                                    max(0, bbox[1] - padding),
                                    min(im.width, bbox[2] + padding),
                                    min(im.height, bbox[3] + padding)
                                )
                                cropped = im.crop(padded_bbox)
                                cropped_dims = (cropped.width, cropped.height)
                                tf = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
                                cropped.save(tf.name)
                                temp_files_to_clean.append(tf.name)
                                image_to_insert = tf.name
                        except Exception as e:
                            print(f"[Doc] 图片裁剪失败 {local}: {e}")

                    p = doc.add_paragraph()

                    if preprocess_images and seq > text_q_seen and not is_diagram:
                        # 序号单独一行，与文字题保持一致
                        p.add_run(f"{seq}、")
                        text_q_seen = seq
                        p.paragraph_format.space_after = Pt(2)

                        p2 = doc.add_paragraph()
                        p2.paragraph_format.space_after = Pt(6)
                        r2 = p2.add_run()
                        if image_to_insert != local and cropped_dims:
                            # 宽度优先缩放：目标满行宽 14cm，按比例算高，限制高度 2.5~6cm
                            aspect = cropped_dims[0] / cropped_dims[1]
                            cm_width = 14.0
                            cm_height = cm_width / aspect
                            if cm_height > 6.0:
                                cm_height = 6.0
                                cm_width = cm_height * aspect
                            if cm_height < 2.5:
                                cm_height = 2.5
                                cm_width = cm_height * aspect
                            if cm_width > 14.0:
                                cm_width = 14.0
                                cm_height = cm_width / aspect
                            r2.add_picture(image_to_insert, width=Cm(cm_width), height=Cm(cm_height))
                        else:
                            shape = r2.add_picture(image_to_insert)
                            _downscale_to_max(shape, EXPORT_IMAGE_MAX_WIDTH, EXPORT_IMAGE_MAX_HEIGHT)
                    else:
                        if is_diagram:
                            if seq > text_q_seen and not seq_has_text.get(seq, False):
                                p.add_run(f"{seq}、")
                                text_q_seen = seq
                            shape = p.add_run().add_picture(image_to_insert)
                            _downscale_to_max(shape, DIAGRAM_IMAGE_MAX_WIDTH, DIAGRAM_IMAGE_MAX_HEIGHT)
                        else:
                            shape = p.add_run().add_picture(image_to_insert)
                            _downscale_to_max(shape, EXPORT_IMAGE_MAX_WIDTH, EXPORT_IMAGE_MAX_HEIGHT)
                            
                    for tf in temp_files_to_clean:
                        try:
                            os.remove(tf)
                        except Exception:
                            pass

        doc.add_page_break()

    # 保存文件
    if not output_path:
        date_list = [str(t.get("date_str") or "")[:10] for t in threads if t.get("date_str")]
        if date_list:
            start_day = min(date_list).replace("-", "")
            end_mmdd = max(date_list)[5:7] + max(date_list)[8:10]
        else:
            start_day = start_date.replace("-", "")
            end_mmdd = end_date[5:7] + end_date[8:10] if len(end_date) >= 10 else end_date.replace("-", "")
        filename = f"{start_day}-{end_mmdd}打卡题.docx"
        output_path = os.path.join(DATA_DIR, filename)

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    print(f"[Doc] 已生成：{output_path}")
    return output_path


def _setup_page(doc: Document):
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)


def _docx_to_pdf(docx_path: str, pdf_path: str) -> str:
    """Convert DOCX to PDF using Pandoc + Chrome headless."""
    import subprocess
    import shutil
    import tempfile
    import sys
    import os

    print("[Doc] 使用 Pandoc + Chrome 转换 (样式可能不完全一致)...")
    html_dir = tempfile.mkdtemp(prefix="shuati_doc_")
    html_path = os.path.join(html_dir, "output.html")

    # Convert DOCX to HTML using pandoc
    subprocess.run(
        ["pandoc", "--standalone", docx_path, "-o", html_path, "--extract-media=" + html_dir],
        check=True,
        capture_output=True,
    )

    # 注入 CSS 模拟 Word 样式
    try:
        with open(html_path, "r", encoding="utf-8") as f:
            html_content = f.read()
        
        css = """
        <style>
        @page {
            size: A4;
            margin: 2cm;
        }
        body {
            font-family: "Songti SC", "SimSun", "Times New Roman", serif;
            font-size: 13pt;
            line-height: 1.5;
            color: #000;
            margin: 0;
            padding: 0;
            max-width: none;
        }
        h1 {
            font-size: 24pt;
            text-align: center;
            color: #000;
            page-break-before: always;
            break-before: page;
            margin-top: 0;
            margin-bottom: 24pt;
        }
        h1:first-of-type {
            page-break-before: auto;
            break-before: auto;
        }
        p { margin-bottom: 6pt; text-align: justify; }
        img { max-width: 100%; height: auto; vertical-align: middle; }
        .center { text-align: center; }
        </style>
        """
        html_content = html_content.replace("</head>", f"{css}</head>")
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(html_content)
    except Exception as e:
        print(f"[Doc] 注入 CSS 失败: {e}")

    # Convert HTML to PDF using Chrome headless
    chrome_paths = [
        "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
        "/opt/homebrew/bin/chromium", 
        "/usr/bin/google-chrome"
    ]
    chrome = next((p for p in chrome_paths if os.path.exists(p)), chrome_paths[0])
    
    abs_pdf_path = os.path.abspath(pdf_path)
    subprocess.run(
        [
            chrome, 
            "--headless", 
            "--disable-gpu", 
            "--no-pdf-header-footer", 
            f"--print-to-pdf={abs_pdf_path}", 
            html_path
        ],
        check=True,
        capture_output=True,
        cwd=html_dir,
    )

    # Cleanup temp dir
    shutil.rmtree(html_dir)

    print(f"[Doc] 已转换：{pdf_path}")
    return pdf_path


def generate_pdf(start_date: str, end_date: str, output_path: str = None, source_thread_ids: list[str] | None = None, preprocess_images: bool = True) -> str:
    """Generate PDF document from questions."""
    final_pdf_path = output_path
    if output_path and output_path.lower().endswith(".pdf"):
        docx_output_path = output_path[:-4] + ".docx"
    else:
        docx_output_path = output_path
        final_pdf_path = None

    docx_path = generate_word(start_date, end_date, docx_output_path, source_thread_ids, preprocess_images)

    if final_pdf_path is None:
        final_pdf_path = docx_path.replace(".docx", ".pdf")

    _docx_to_pdf(docx_path, final_pdf_path)
    return final_pdf_path
